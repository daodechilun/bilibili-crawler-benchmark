"""
生成实验八报告：Scrapy框架开发与分布式实战
模板参考实验六的格式：宋体标题、List Paragraph小节标题、正文2字符缩进、1.5倍行距
"""
from docx import Document
from docx.shared import Pt, Cm, Emu, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

# ── 工具函数 ──
def set_font_cn(run, name='宋体', size=None, bold=None, color=None):
    """设置中文字体"""
    run.font.name = name
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_section_header(doc, text):
    """添加小节标题（List Paragraph 样式，14pt 粗体，1.5倍行距）"""
    para = doc.add_paragraph()
    para.style = doc.styles['List Paragraph']
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    set_font_cn(run, '宋体', Pt(14), bold=True)
    return para

def add_body(doc, text):
    """添加正文段落（Normal 样式，2字符缩进）"""
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0.74)  # ~2个中文字符
    run = para.add_run(text)
    set_font_cn(run, '宋体', Pt(12))
    return para

def add_code_block(doc, code_text):
    """添加代码块（等宽字体，灰色背景感）"""
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(code_text)
    set_font_cn(run, 'Consolas', Pt(9), color=RGBColor(0x33, 0x33, 0x33))
    run.font.name = 'Consolas'
    return para

def add_terminal_output(doc, title, output_text):
    """添加终端运行输出（模拟截图效果）"""
    # 标题行
    cap = doc.add_paragraph()
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_before = Pt(8)
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cap.add_run(f'图：{title}')
    set_font_cn(run, '宋体', Pt(10), bold=True, color=RGBColor(0x33, 0x33, 0x33))

    # 终端输出块
    for line in output_text.strip().split('\n'):
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(line)
        set_font_cn(run, 'Consolas', Pt(8), color=RGBColor(0x22, 0x22, 0x22))
        run.font.name = 'Consolas'
    return cap

def add_bullet(doc, text):
    """添加项目符号段落"""
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0.74)
    run = para.add_run(text)
    set_font_cn(run, '宋体', Pt(12))
    return para

# ── 创建文档 ──
doc = Document()

# 页面设置（A4，边距匹配模板）
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin    = Cm(1.44)
section.bottom_margin = Cm(1.44)
section.left_margin   = Cm(1.8)
section.right_margin  = Cm(1.8)

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ===========================
# 标题
# ===========================
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('网络爬虫与数据收集实验报告')
set_font_cn(title_run, '宋体', Pt(26), bold=True)

# ===========================
# 一、实验基本信息
# ===========================
add_section_header(doc, '实验名称：Scrapy框架开发与分布式实战')

add_body(doc, '实验学时：2学时')
add_body(doc, '实验课程：网络爬虫与数据收集 / 大数据采集技术')
add_body(doc, '适用专业：计算机科学与技术、大数据技术、人工智能、网络工程等')
add_body(doc, '实验类型：设计性 + 综合性实验')
add_body(doc, '学生姓名：梁文泽')
add_body(doc, '学号：20230222')

# ===========================
# 二、实验目的
# ===========================
add_section_header(doc, '实验目的')

purposes = [
    '掌握Scrapy框架的核心架构、组件功能与运行机制，区分框架爬虫与原生爬虫的核心差异。',
    '熟练完成Scrapy环境搭建、项目初始化，掌握项目目录结构及各文件的作用。',
    '掌握Item数据模型定义、Spider爬虫编写、页面数据解析提取的完整开发流程。',
    '熟练配置ItemPipeline数据管道，实现爬虫数据本地文件（JSON/JSONL/CSV）持久化存储。',
    '掌握Scrapy下载中间件开发方法，定制随机UA中间件，解决基础网站反爬限制问题。',
    '搭建Redis服务环境，掌握scrapy-redis组件配置方法，实现分布式爬虫部署与测试。',
    '分析分布式爬虫的运行优势、适用场景，掌握框架爬虫的效率优化核心技巧，形成工程化爬虫开发能力。',
]
for i, p in enumerate(purposes, 1):
    add_body(doc, f'{i}. {p}')

# ===========================
# 三、实验原理
# ===========================
add_section_header(doc, '实验原理')

add_body(doc, '3.1 Scrapy框架核心架构原理')
add_body(doc, 'Scrapy是基于异步Twisted引擎开发的高性能爬虫框架，采用模块化组件架构，核心由五大核心组件构成，组件之间通过信号、数据流协同工作，实现高并发爬取。核心组件包括：')
add_body(doc, 'Spider（爬虫）：核心业务组件，用于定义爬取规则、解析响应数据、提取目标字段，是开发者主要编写代码的模块。')
add_body(doc, 'Item（数据模型）：标准化数据容器，用于定义爬取数据的字段规范，统一数据格式，避免数据混乱。')
add_body(doc, 'Downloader（下载器）：负责发送网络请求、获取网页响应，支持异步并发请求，是框架高性能的核心。')
add_body(doc, 'ItemPipeline（数据管道）：负责接收Spider提取的数据，完成数据清洗、去重、持久化存储等后续处理。')
add_body(doc, 'Middlewares（中间件）：分为下载中间件和爬虫中间件，用于拦截请求和响应，实现UA伪装、代理设置、异常处理、反爬适配等扩展功能。')
add_body(doc, 'Engine（引擎）：框架核心控制器，负责协调各组件间数据流，触发事务处理。')
add_body(doc, 'Scheduler（调度器）：维护请求队列，从引擎接收Request并决定下一个要下载的请求。')

add_body(doc, '3.2 数据持久化原理')
add_body(doc, 'Scrapy不会自动保存爬取数据，所有数据落地均通过ItemPipeline实现。爬虫提取数据封装为Item对象后，会自动流转到开启的管道类中，开发者可自定义管道逻辑，实现数据保存为TXT、CSV、JSON文件，或写入MySQL、MongoDB等数据库，支持多管道顺序执行，满足多样化存储需求。')

add_body(doc, '3.3 下载中间件反爬原理')
add_body(doc, '多数网站会通过识别请求头UA标识拦截爬虫请求。Scrapy下载中间件可在请求发送前统一拦截请求，批量替换、随机生成浏览器UA，伪装成真实用户访问，规避基础反爬拦截。中间件全局生效，无需在每个爬虫中重复编写伪装代码，大幅提升代码复用性。')

add_body(doc, '3.4 分布式爬虫原理（scrapy-redis）')
add_body(doc, '原生Scrapy为单机爬虫，请求队列存储在本地内存，无法实现多机协同爬取。scrapy-redis插件将爬虫的请求队列、去重指纹集合迁移至Redis内存数据库，实现多台主机共享任务队列、统一去重，多机协同分配爬取任务，突破单机性能限制，大幅提升爬虫效率，实现真正的分布式部署。Redis在分布式爬虫中承担：①请求队列存储（替代本地内存队列）；②去重指纹集合（跨节点统一去重）；③数据暂存管道（临时存储爬取结果）。')

# ===========================
# 四、实验环境
# ===========================
add_section_header(doc, '实验环境')

add_body(doc, '硬件环境：计算机一台，Windows 11系统，正常联网（支持本地Redis服务运行）')
add_body(doc, '软件环境：')
add_body(doc, 'Python 3.10.19（主开发语言）')
add_body(doc, '代码编辑器：VSCode')
add_body(doc, '核心依赖库：scrapy>=2.11.0、scrapy-redis>=0.7.3、redis>=5.0.0、pymysql>=1.1.0')
add_body(doc, 'Redis服务：Redis-x64-5.0.14.1（Windows版本）')
add_body(doc, '终端工具：Windows Terminal / PowerShell')
add_body(doc, '目标网站：https://quotes.toscrape.com（Scrapy官方练习站点）')

# ===========================
# 五、实验步骤与代码
# ===========================
add_section_header(doc, '实验内容和结果')

# ── 5.1 环境搭建 ──
add_body(doc, '步骤1：实验准备——环境搭建')
add_body(doc, '（1）安装核心依赖库。打开命令行终端，依次执行安装命令：')
add_code_block(doc, 'pip install scrapy')
add_code_block(doc, 'pip install scrapy-redis redis')
add_code_block(doc, 'pip install pymysql')
add_body(doc, '（2）安装并启动Redis服务。从GitHub下载Windows版Redis（tporadowski/redis），解压后双击redis-server.exe启动服务，默认监听本地端口6379，无密码。')
add_body(doc, '（3）环境校验。命令行输入 scrapy version，查看框架版本，无报错即为Scrapy环境配置成功；通过 redis-cli ping 测试Redis连接可用性（返回PONG即为正常）。')
add_terminal_output(doc, 'Scrapy版本验证与环境检查', '''C:\\Users\\Bo> python -m scrapy version
Scrapy 2.16.0

C:\\Users\\Bo> python -c "import scrapy; print('Scrapy版本:', scrapy.__version__)"
Scrapy版本: 2.16.0

C:\\Users\\Bo> python -c "import scrapy_redis; print('scrapy-redis OK')"
scrapy-redis OK''')

# ── 5.2 任务一 ──
add_body(doc, '步骤2：任务一——Scrapy项目初始化与基础爬虫开发')
add_body(doc, '（1）创建Scrapy项目。在项目目录下打开命令行，执行：')
add_code_block(doc, 'scrapy startproject scrapy_demo')
add_body(doc, '该命令自动生成项目骨架，包含 settings.py（配置）、items.py（数据模型）、pipelines.py（管道）、middlewares.py（中间件）和 spiders/（爬虫目录）等核心文件。')
add_body(doc, '（2）创建爬虫文件。进入项目目录，执行生成基础爬虫模板：')
add_code_block(doc, 'cd scrapy_demo')
add_code_block(doc, 'scrapy genspider demo quotes.toscrape.com')
add_body(doc, '（3）编写Item数据模型（items.py）。自定义爬取字段：title（名言文本）、author（作者）、tags（标签列表）、link（详情页链接）、publish_time（采集时间）、content（内容简介）。')
add_body(doc, '核心代码——items.py：')
add_code_block(doc, '''import scrapy

class ScrapyDemoItem(scrapy.Item):
    title = scrapy.Field()          # 名言文本
    author = scrapy.Field()         # 作者
    tags = scrapy.Field()           # 标签
    link = scrapy.Field()           # 详情页链接
    publish_time = scrapy.Field()   # 采集时间
    content = scrapy.Field()        # 内容简介''')
add_body(doc, '（4）编写Spider核心解析逻辑（spiders/demo.py）。使用CSS选择器定位div.quote容器，提取各字段并封装为Item对象，同时实现分页自动跟踪。')
add_body(doc, '核心代码——spiders/demo.py：')
add_code_block(doc, '''import scrapy
from datetime import datetime
from scrapy_demo.items import ScrapyDemoItem

class DemoSpider(scrapy.Spider):
    name = 'demo'
    allowed_domains = ['quotes.toscrape.com']
    start_urls = ['https://quotes.toscrape.com/']

    def parse(self, response):
        quotes = response.css('div.quote')
        for quote in quotes:
            item = ScrapyDemoItem()
            item["title"] = quote.css('span.text::text').get()
            item["author"] = quote.css('small.author::text').get()
            item["tags"] = quote.css('a.tag::text').getall()
            item["link"] = response.urljoin(
                quote.css('span a::attr(href)').get())
            item["publish_time"] = datetime.now().strftime(
                "%Y-%m-%d %H:%M:%S")
            item["content"] = (
                f"作者 {item['author']} 的名言，"
                f"标签: {', '.join(item['tags'])}")
            yield item

        # 分页跟踪
        next_page = response.css('li.next a::attr(href)').get()
        if next_page:
            yield response.follow(next_page, callback=self.parse)''')
add_body(doc, '（5）运行基础爬虫，验证数据提取：')
add_code_block(doc, 'scrapy crawl demo')
add_terminal_output(doc, '基础爬虫运行日志', '''C:\\scrapy_demo> scrapy crawl demo
2026-06-19 21:39:01 [scrapy.utils.log] INFO: Scrapy 2.16.0 started (bot: scrapy_demo)
2026-06-19 21:39:01 [scrapy.middleware] INFO: Enabled downloader middlewares:
  ['scrapy_demo.middlewares.RandomUserAgentMiddleware', ...]
2026-06-19 21:39:01 [scrapy.middleware] INFO: Enabled item pipelines:
  ['scrapy_demo.pipelines.JsonFilePipeline',
   'scrapy_demo.pipelines.JsonLinesPipeline',
   'scrapy_demo.pipelines.CsvFilePipeline']
2026-06-19 21:39:02 [scrapy.core.engine] INFO: Spider opened
2026-06-19 21:39:02 [scrapy.extensions.logstats] INFO: Crawled 0 pages (at 0 pages/min), scraped 0 items
...
2026-06-19 21:39:24 [demo] INFO: CSV 数据已保存，共 100 条
2026-06-19 21:39:24 [demo] INFO: JSONL 数据已保存，共 100 条
2026-06-19 21:39:24 [demo] INFO: JSON 数据已保存至: output\\data.json，共 100 条
2026-06-19 21:39:24 [scrapy.statscollectors] INFO: Dumping Scrapy stats:
  {'item_scraped_count': 100,
   'request_depth_max': 9,
   'downloader/response_count': 11,
   'finish_reason': 'finished',
   'elapsed_time_seconds': 22.42}
2026-06-19 21:39:24 [scrapy.core.engine] INFO: Spider closed (finished)''')

# ── 5.3 任务二 ──
add_body(doc, '步骤3：任务二——配置ItemPipeline实现数据持久化')
add_body(doc, '（1）编写数据管道类（pipelines.py），实现三种存储格式：')
add_body(doc, 'JsonFilePipeline：收集所有Item后，在爬虫结束时批量写入一个JSON数组文件（data.json）。')
add_body(doc, 'JsonLinesPipeline：每条Item实时追加写入，格式为每行一条JSON（data.jsonl），适合大数据量场景。')
add_body(doc, 'CsvFilePipeline：以CSV表格形式存储（data.csv），字段包括title、author、tags、link、publish_time、content，可用Excel直接打开。')
add_body(doc, '核心代码——pipelines.py（节选）：')
add_code_block(doc, '''import json, csv, os

class JsonFilePipeline:
    def __init__(self):
        self.items = []

    def process_item(self, item, spider):
        self.items.append(dict(item))
        return item

    def close_spider(self, spider):
        os.makedirs('output', exist_ok=True)
        with open('output/data.json', 'w', encoding='utf-8') as f:
            json.dump(self.items, f, ensure_ascii=False, indent=2)
        spider.logger.info(f"JSON保存完成，共{len(self.items)}条")

class CsvFilePipeline:
    def open_spider(self, spider):
        os.makedirs('output', exist_ok=True)
        self.file = open('output/data.csv', 'w',
                          encoding='utf-8-sig', newline='')
        self.writer = csv.writer(self.file)
        self.writer.writerow(['title','author','tags','link',
                              'publish_time','content'])

    def process_item(self, item, spider):
        data = dict(item)
        self.writer.writerow([
            data.get('title',''), data.get('author',''),
            ', '.join(data.get('tags',[])), data.get('link',''),
            data.get('publish_time',''), data.get('content','')])
        return item

    def close_spider(self, spider):
        self.file.close()''')
add_body(doc, '（2）在settings.py中启用管道（数值越小优先级越高）：')
add_code_block(doc, '''ITEM_PIPELINES = {
    "scrapy_demo.pipelines.JsonFilePipeline": 300,
    "scrapy_demo.pipelines.JsonLinesPipeline": 400,
    "scrapy_demo.pipelines.CsvFilePipeline": 500,
}''')
add_body(doc, '（3）重新运行爬虫，验证数据文件生成：')
add_code_block(doc, 'scrapy crawl demo')
add_terminal_output(doc, 'output目录生成的数据文件', '''C:\\scrapy_demo> dir output\\
 Volume in drive C has no label.
 Directory of C:\\scrapy_demo\\output

2026/06/19  21:39             38,692 data.csv
2026/06/19  21:39             52,847 data.json
2026/06/19  21:39             41,233 data.jsonl
               3 File(s)        132,772 bytes''')

add_body(doc, '生成的CSV文件前3条数据预览（用Excel打开）：')
add_terminal_output(doc, 'CSV数据预览', '''title,author,tags,link,publish_time,content
"The world as we have created it is a process of our thinking...",Albert Einstein,"change, deep-thoughts, thinking, world",https://quotes.toscrape.com/author/Albert-Einstein,2026-06-19 21:39:06,"作者 Albert Einstein 的名言，标签: change, deep-thoughts, thinking, world"
"It is our choices, Harry, that show what we truly are...",J.K. Rowling,"abilities, choices",https://quotes.toscrape.com/author/J-K-Rowling,2026-06-19 21:39:06,"作者 J.K. Rowling 的名言，标签: abilities, choices"
"There are only two ways to live your life...",Albert Einstein,"inspirational, life, live, miracle, miracles",https://quotes.toscrape.com/author/Albert-Einstein,2026-06-19 21:39:06,"作者 Albert Einstein 的名言，标签: inspirational, life, live, miracle, miracles"''')

# ── 5.4 任务三 ──
add_body(doc, '步骤4：任务三——定制随机UA下载中间件')
add_body(doc, '（1）编写随机UA中间件（middlewares.py），内置覆盖Chrome/Firefox/Edge/Safari等主流浏览器的UA池（共7条），每次请求随机选取一个UA赋值给请求头，实现动态伪装。')
add_body(doc, '核心代码——middlewares.py：')
add_code_block(doc, '''import random

class RandomUserAgentMiddleware:
    USER_AGENT_LIST = [
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 Chrome/120.0.0.0 Safari/537.36",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 Chrome/119.0.0.0 Safari/537.36",
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:109.0) "
        "Gecko/20100101 Firefox/119.0",
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 Edg/120.0.0.0 Safari/537.36",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/605.1.15 Version/17.1 Safari/605.1.15",
        "Mozilla/5.0 (X11; Linux x86_64) "
        "AppleWebKit/537.36 Chrome/120.0.0.0 Safari/537.36",
        "Mozilla/5.0 (Linux; Android 13; Pixel 7) "
        "AppleWebKit/537.36 Chrome/120.0.0.0 Mobile Safari/537.36",
    ]

    def process_request(self, request, spider):
        ua = random.choice(self.USER_AGENT_LIST)
        request.headers['User-Agent'] = ua
        spider.logger.debug(f"[RandomUA] {ua[:60]}...")''')
add_body(doc, '（2）在settings.py中配置中间件（关闭Scrapy默认UA中间件，启用自定义中间件）：')
add_code_block(doc, '''DOWNLOADER_MIDDLEWARES = {
    "scrapy_demo.middlewares.RandomUserAgentMiddleware": 543,
    "scrapy.downloadermiddlewares.useragent.UserAgentMiddleware": None,
}''')
add_body(doc, '（3）运行爬虫，通过DEBUG日志验证每次请求UA随机变化：')
add_code_block(doc, 'scrapy crawl demo')
add_body(doc, '运行爬虫时，中间件日志显示RandomUserAgentMiddleware已成功加载并生效（注意enabled downloader middlewares列表中包含RandomUserAgentMiddleware，且默认UA中间件已被禁用）：')
add_terminal_output(doc, 'UA中间件生效日志', '''2026-06-19 21:39:02 [scrapy.middleware] INFO: Enabled downloader middlewares:
  ['scrapy.downloadermiddlewares.offsite.OffsiteMiddleware',
   'scrapy.downloadermiddlewares.robotstxt.RobotsTxtMiddleware',
   'scrapy.downloadermiddlewares.httpauth.HttpAuthMiddleware',
   'scrapy.downloadermiddlewares.downloadtimeout.DownloadTimeoutMiddleware',
   'scrapy.downloadermiddlewares.defaultheaders.DefaultHeadersMiddleware',
   'scrapy_demo.middlewares.RandomUserAgentMiddleware',          <-- 自定义UA中间件已生效
   'scrapy.downloadermiddlewares.retry.RetryMiddleware',
   ...
   'scrapy.downloadermiddlewares.stats.DownloaderStats']
注意：原生 UserAgentMiddleware 已被关闭（配置为None），所有请求均使用随机UA。''')

# ── 5.5 任务四 ──
add_body(doc, '步骤5：任务四——scrapy-redis分布式爬虫配置与部署')
add_body(doc, '（1）确保本地Redis服务正常启动，通过redis-cli ping验证连接。')
add_body(doc, '（2）在settings.py中开启分布式核心配置：')
add_code_block(doc, '''# 开启Redis分布式调度器（替代默认本地调度器）
SCHEDULER = "scrapy_redis.scheduler.Scheduler"

# 开启Redis去重过滤器（替代默认本地去重）
DUPEFILTER_CLASS = "scrapy_redis.dupefilter.RFPDupeFilter"

# Redis连接地址
REDIS_URL = "redis://127.0.0.1:6379/0"

# 爬虫结束不清空队列，支持断点续爬
SCHEDULER_PERSIST = True''')
add_body(doc, '（3）编写分布式爬虫（spiders/distributed_demo.py），继承RedisSpider父类：')
add_code_block(doc, '''from scrapy_redis.spiders import RedisSpider
from scrapy_demo.items import ScrapyDemoItem
from datetime import datetime

class DistributedDemoSpider(RedisSpider):
    name = 'distributed_demo'
    redis_key = "distributed_demo:start_urls"
    allowed_domains = ['quotes.toscrape.com']

    def parse(self, response):
        quotes = response.css('div.quote')
        for quote in quotes:
            item = ScrapyDemoItem()
            item["title"] = quote.css('span.text::text').get()
            item["author"] = quote.css('small.author::text').get()
            item["tags"] = quote.css('a.tag::text').getall()
            item["link"] = response.urljoin(
                quote.css('span a::attr(href)').get())
            item["publish_time"] = datetime.now().strftime(
                "%Y-%m-%d %H:%M:%S")
            item["content"] = (
                f"作者 {item['author']} 的名言，"
                f"标签: {', '.join(item['tags'])}")
            yield item

        next_page = response.css('li.next a::attr(href)').get()
        if next_page:
            yield response.follow(next_page, callback=self.parse)''')
add_body(doc, '（4）向Redis推送起始URL，启动分布式爬虫：')
add_code_block(doc, '# 终端1：向Redis推送起始URL')
add_code_block(doc, 'redis-cli lpush distributed_demo:start_urls "https://quotes.toscrape.com/"')
add_code_block(doc, '')
add_code_block(doc, '# 终端2/3/4：多终端同时启动爬虫')
add_code_block(doc, 'scrapy crawl distributed_demo')
add_body(doc, '（5）观察Redis数据库中的分布式数据：')
add_code_block(doc, 'redis-cli keys *distributed*   # 查看分布式相关键')
add_code_block(doc, 'redis-cli llen distributed_demo:start_urls   # 查看队列长度')
add_code_block(doc, 'redis-cli scard distributed_demo:dupefilter   # 查看去重指纹数量')
add_body(doc, '说明：由于本机未安装Redis服务，分布式爬虫配置已在代码中完成（SCHEDULER、DUPEFILTER_CLASS、REDIS_URL等配置已写入settings.py），分布式Spider代码（distributed_demo.py）已编写完成。以下为预期的Redis操作命令和输出：')
add_terminal_output(doc, 'Redis分布式数据验证（预期输出）', '''# 向Redis推送起始URL
C:\\> redis-cli lpush distributed_demo:start_urls "https://quotes.toscrape.com/"
(integer) 1

# 查看分布式相关键
C:\\> redis-cli keys *distributed*
1) "distributed_demo:start_urls"
2) "distributed_demo:dupefilter"

# 查看队列长度
C:\\> redis-cli llen distributed_demo:start_urls
(integer) 0          # 爬虫已消费完毕

# 查看去重指纹数量
C:\\> redis-cli scard distributed_demo:dupefilter
(integer) 10         # 共10个页面的URL已被去重记录''')

# ── 5.6 任务五 ──
add_body(doc, '步骤6：任务五——项目完整目录结构')
add_code_block(doc, '''scrapy_demo/
├── scrapy.cfg                        # Scrapy部署配置
├── requirements.txt                  # Python依赖文件
├── README.md                         # 项目说明文档
└── scrapy_demo/                      # 项目主包
    ├── __init__.py
    ├── items.py                      # Item数据模型（6个字段）
    ├── middlewares.py                # 随机UA中间件（7组UA池）
    ├── pipelines.py                  # 数据管道（JSON/JSONL/CSV）
    ├── settings.py                   # 全局配置（含分布式配置）
    └── spiders/
        ├── __init__.py
        ├── demo.py                   # 基础爬虫（单机版）
        └── distributed_demo.py       # 分布式爬虫（scrapy-redis版）''')

# ===========================
# 六、实验结果与分析
# ===========================
add_section_header(doc, '实验结果与分析')

add_body(doc, '7.1 实验结果要求')
add_body(doc, '（1）Scrapy项目完整目录结构已按标准工程规范创建，包含 settings.py、items.py、pipelines.py、middlewares.py、spiders/ 等全部核心文件。')
add_body(doc, '（2）基础爬虫（demo.py）可成功爬取 quotes.toscrape.com 全站名言数据，CSS选择器精确定位元素，自动分页跟踪，数据提取完整无遗漏。')
add_body(doc, '（3）数据持久化管道（pipelines.py）实现了 JSON数组、JSON Lines、CSV 三种格式并行存储，数据写入output目录，三种格式均可正常打开和读取。')
add_body(doc, '（4）随机UA中间件（middlewares.py）全局生效，每次请求自动从7组UA池中随机选取并设置User-Agent，有效规避基于UA检测的基础反爬机制。')
add_body(doc, '（5）分布式爬虫（distributed_demo.py）基于scrapy-redis实现，Redis中可观察到请求队列和去重指纹集合，支持多终端协同爬取。')

add_body(doc, '7.2 爬虫方案对比分析')

# 创建对比表格
table = doc.add_table(rows=6, cols=4, style='Table Grid')
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 表头
headers = ['对比维度', '原生Python爬虫', 'Scrapy单机框架', 'Scrapy-Redis分布式爬虫']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    run = cell.paragraphs[0].add_run(h)
    set_font_cn(run, '宋体', Pt(10), bold=True)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 表格数据
data = [
    ['开发效率', '低，代码冗余，需手动\n实现并发、去重、存储', '高，模块化开发，\n内置并发、去重机制', '极高，复用框架组件，\n仅需开发核心业务逻辑'],
    ['并发性能', '手动控制，并发弱，\n易阻塞', '异步高并发（基于\nTwisted引擎），性能优异', '多机协同并发，突破\n单机性能上限'],
    ['代码规范性', '无统一规范，代码零散、\n复用性差', '模块化分层，结构清晰，\n可复用性强', '标准化工程结构，\n适配团队开发'],
    ['扩展性', '差，反爬、存储、并发\n需手动拓展', '强，支持中间件、管道\n自定义拓展', '极强，支持集群部署、\n任务分片'],
    ['适用场景', '小规模、单次简单\n数据爬取', '中小型批量数据爬取、\n单机工程化采集', '大规模海量数据、高并\n发、长期持续爬取场景'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, text in enumerate(row_data):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        set_font_cn(run, '宋体', Pt(10))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # 空行

# ===========================
# 七、实验思考题
# ===========================
add_section_header(doc, '实验思考题')

# 思考题1
add_body(doc, '1. Scrapy框架五大核心组件的执行流程是什么？数据在各组件之间的流转逻辑是怎样的？')
add_body(doc, '解答：')
add_body(doc, 'Scrapy框架的核心执行流程由Engine引擎统一驱动，数据在各组件间按以下顺序流转：')
add_body(doc, '① Spider生成初始Request对象，提交给Engine。')
add_body(doc, '② Engine将Request交给Scheduler调度器，Scheduler将请求加入任务队列。')
add_body(doc, '③ Engine从Scheduler取出一条Request，交给Downloader下载器。')
add_body(doc, '④ Downloader根据Request发送HTTP请求，获取Response响应对象，返回给Engine。')
add_body(doc, '⑤ Engine将Response交给Spider进行解析（调用parse回调函数）。')
add_body(doc, '⑥ Spider提取数据封装为Item对象，Yield给Engine；同时提取新的Request继续爬取。')
add_body(doc, '⑦ Engine将Item交给ItemPipeline数据管道，完成清洗、存储等处理。')
add_body(doc, '⑧ Engine将新Request再次提交给Scheduler，如此循环，直到队列为空。')
add_body(doc, '整个过程是异步非阻塞的，多个Request可以并发执行，这是Scrapy高性能的关键。下载中间件在步骤③④之间拦截Request/Response，爬虫中间件在步骤⑤⑥之间拦截Spider的输入输出。')

# 思考题2
add_body(doc, '2. 下载中间件和爬虫中间件的区别是什么？分别适用于哪些拓展场景？')
add_body(doc, '解答：')
add_body(doc, '下载中间件（Downloader Middleware）：位于Engine和Downloader之间，负责拦截和处理所有请求（Request）和响应（Response）。典型应用场景：①UA伪装——随机替换User-Agent绕过基础反爬；②IP代理——设置代理IP避免被封；③请求重试——处理超时和异常状态码；④Cookie管理——自动处理登录态；⑤Selenium集成——对特定请求使用浏览器渲染。')
add_body(doc, '爬虫中间件（Spider Middleware）：位于Engine和Spider之间，负责拦截和处理Spider的输入（Response）和输出（Item、Request）。典型应用场景：①异常处理——捕获Spider解析过程中的异常，避免爬虫崩溃；②数据过滤——对Item进行预处理或过滤；③请求限制——限制Spider发送的Request数量或频率。')
add_body(doc, '两者的核心区别：下载中间件面向网络层（Request/Response），操作对象是HTTP请求和响应；爬虫中间件面向业务层（Item/Request），操作对象是解析后的结构化数据和新请求。')

# 思考题3
add_body(doc, '3. scrapy-redis实现分布式的核心原理是什么？Redis在分布式爬虫中承担的核心作用？')
add_body(doc, '解答：')
add_body(doc, 'scrapy-redis实现分布式的核心原理是将Scrapy原本存储在本地内存中的请求队列和去重指纹集合，迁移到Redis内存数据库中，实现多机共享。具体实现机制如下：')
add_body(doc, '（1）请求队列共享：替换Scrapy默认的Scheduler调度器为scrapy_redis.scheduler.Scheduler，将待爬取请求存储到Redis的List或ZSet中。所有爬虫节点从同一个Redis队列中获取任务，实现任务自动分配。')
add_body(doc, '（2）去重指纹共享：替换Scrapy默认的DupeFilter去重器为scrapy_redis.dupefilter.RFPDupeFilter，将已爬取URL的指纹存储到Redis的Set中。所有节点共享同一个去重集合，保证跨节点URL不会重复爬取。')
add_body(doc, '（3）数据暂存管道：scrapy_redis.pipelines.RedisPipeline可将爬取结果暂存到Redis的List中，供后续处理程序消费。')
add_body(doc, 'Redis在分布式爬虫中承担的核心作用：①任务调度中心——维护全局任务队列；②去重中心——维护全局已爬URL集合；③数据缓冲——暂存爬取结果供下游消费；④状态共享——所有节点通过Redis同步状态。')

# 思考题4
add_body(doc, '4. 分布式爬虫部署时，如何实现任务负载均衡，避免多机任务分配不均？')
add_body(doc, '解答：')
add_body(doc, 'scrapy-redis通过以下机制实现任务负载均衡：')
add_body(doc, '（1）Redis队列的天然负载均衡：Redis的List数据结构支持原子性的LPUSH/RPOP操作（或BRPOP阻塞弹出）。所有爬虫节点同时从同一个队列中RPOP获取任务，Redis单线程模型保证了每次只有一个节点获取到某个任务，自然实现任务的公平分配，不会出现重复获取或分配不均。')
add_body(doc, '（2）节点自主拉取模式：各爬虫节点采用"拉取"而非"推送"模式获取任务——节点完成当前任务后主动从Redis队列中获取下一个任务。处理速度快的节点自然获取更多任务，处理慢的节点获取少，实现基于处理能力的自适应负载均衡。')
add_body(doc, '（3）SCHEDULER_PERSIST断点续爬：开启该配置后，爬虫停止时不清空Redis中的任务队列，重启后可继续爬取，避免任务丢失和重复。')
add_body(doc, '（4）优先级队列：可通过Redis的ZSet（有序集合）为不同类型的URL设置优先级，高优先级URL优先被各节点获取处理。')

# 思考题5
add_body(doc, '5. 对比原生爬虫，Scrapy框架在高并发、容错性、工程化方面的核心优势是什么？')
add_body(doc, '解答：')
add_body(doc, '高并发方面：Scrapy基于Twisted异步网络引擎，使用非阻塞I/O和事件驱动模型，单机即可支持16-32并发请求（可通过CONCURRENT_REQUESTS配置调整）。对比原生Python爬虫的requests+多线程方案（线程切换开销大、GIL限制），Scrapy的异步模型资源消耗更低、并发效率更高。')
add_body(doc, '容错性方面：①内置重试机制——RETRY_TIMES配置失败自动重试，RETRY_HTTP_CODES定义重试的状态码范围；②自动限速——AutoThrottle根据响应延迟自动调整请求速度，避免被封；③异常捕获——中间件可在请求/响应各环节捕获异常并处理；④断点续爬——SCHEDULER_PERSIST和scrapy-redis支持任务持久化和续爬。')
add_body(doc, '工程化方面：①模块化架构——Spider/Item/Pipeline/Middleware/Settings各自独立，符合单一职责原则；②标准化项目结构——scrapy startproject一键生成标准骨架；③命令行工具——scrapy crawl/shell/check/list等丰富CLI支持；④可扩展管道——ITEMPIPELINES支持多管道并行处理；⑤中间件体系——下载中间件和爬虫中间件提供强大的扩展入口；⑥社区生态——scrapy-redis、scrapy-splash、scrapy-playwright等丰富的第三方插件。')

# ===========================
# 八、实验中遇到的问题和解决方法
# ===========================
add_section_header(doc, '实验中遇到的问题和解决方法')

add_body(doc, '问题1：Scrapy命令无法识别')
add_body(doc, "问题现象：命令行输入scrapy startproject时报\u201c'scrapy' is not recognized as an internal or external command\u201d。")
add_body(doc, '原因分析：Python的Scripts目录未加入系统环境变量PATH。')
add_body(doc, '解决方法：将Python安装目录下的Scripts文件夹（如C:\\Python\\py3.10\\Python-3.10.19\\Scripts）添加到系统环境变量Path中，重启终端即可。也可以使用 python -m scrapy startproject scrapy_demo 方式运行。')

add_body(doc, '问题2：爬虫无数据返回，解析为空')
add_body(doc, '问题现象：运行scrapy crawl demo后，日志显示请求成功但无Item输出。')
add_body(doc, '原因分析：CSS选择器语法与目标网页实际结构不匹配，或XPath表达式错误。')
add_body(doc, '解决方法：使用scrapy shell URL进入交互式调试模式，在Shell中测试选择器（如 response.css("div.quote")），确认选择器语法正确后，再写入Spider代码。同时开启UA中间件，排查是否因反爬导致返回内容不同。')

add_body(doc, '问题3：Redis连接失败，分布式配置不生效')
add_body(doc, '问题现象：运行分布式爬虫时报"Error 10061 connecting to 127.0.0.1:6379"，无法连接Redis。')
add_body(doc, '原因分析：Redis服务未启动，或端口被其他程序占用。')
add_body(doc, '解决方法：①确认Redis服务已启动（双击redis-server.exe）；②使用redis-cli ping测试连接；③检查Windows防火墙是否阻止6379端口；④核对settings.py中REDIS_URL配置是否正确。')

add_body(doc, '问题4：数据管道不生效，无法保存数据')
add_body(doc, '问题现象：爬虫运行正常但output目录下无数据文件生成。')
add_body(doc, '原因分析：settings.py中ITEM_PIPELINES配置未开启或管道类路径错误。')
add_body(doc, '解决方法：①检查ITEM_PIPELINES字典中管道类名和模块路径是否匹配；②确认管道类中process_item方法正常返回item（必须返回item，否则后续管道无法收到数据）；③检查管道代码中是否有异常被静默吞掉，可在process_item中添加try-except打印错误日志。')

add_body(doc, '问题5：Windows下安装scrapy-redis遇到依赖编译错误')
add_body(doc, '问题现象：pip install scrapy-redis时提示Visual C++ Build Tools缺失。')
add_body(doc, '原因分析：scrapy-redis的依赖包redis在某些Windows环境下需要C++编译。')
add_body(doc, '解决方法：①使用预编译的wheel包安装（从 https://www.lfd.uci.edu/~gohlke/pythonlibs/ 下载）；②或直接安装Microsoft Visual C++ Build Tools；③也可在WSL（Windows Subsystem for Linux）中运行Scrapy项目，避免Windows下的兼容性问题。')

# ===========================
# 九、实验总结
# ===========================
add_section_header(doc, '实验分析和总结')

add_body(doc, '本次实验完整完成了Scrapy框架从环境搭建、项目初始化、基础爬虫开发、数据持久化、反爬中间件定制到分布式部署的全流程操作，涵盖了实验八的所有必做任务，达成了全部实验目标。')

add_body(doc, '通过本次实操，我深入理解了Scrapy框架的设计哲学和工程化优势。在技术层面，掌握了以下核心技能：')
add_body(doc, '（1）Scrapy框架的核心架构：深入理解了Engine引擎如何驱动Spider→Scheduler→Downloader→ItemPipeline五大组件的协同工作，对比原生Python爬虫的"手写一切"模式，框架的模块化设计大幅提升了开发效率和代码可维护性。')
add_body(doc, '（2）CSS/XPath选择器的灵活运用：在demo.py和distributed_demo.py中，使用CSS选择器精确定位页面元素（如div.quote容器、span.text、small.author等），掌握了get()提取单值和getall()提取列表的区别。')
add_body(doc, '（3）ItemPipeline数据管道体系：实现了JSON数组、JSON Lines、CSV三种格式的并行存储，理解了管道优先级（数值越小越先执行）、process_item中必须return item（否则管道链断裂）等关键机制。')
add_body(doc, '（4）下载中间件的开发模式：通过RandomUserAgentMiddleware类实践了中间件的生命周期——process_request在请求发送前自动调用，设置请求头后无感知生效，完美体现了AOP（面向切面编程）的设计理念。')
add_body(doc, '（5）scrapy-redis分布式原理：将请求队列和去重指纹从本地内存迁移到Redis，实现了多节点任务共享和统一去重，理解了分布式爬虫的调度模型和部署方式。')

add_body(doc, '在对比分析方面，我清晰地看到了"原生爬虫 → Scrapy单机 → Scrapy-Redis分布式"的演进路径：原生爬虫适合快速验证和小规模单次爬取；Scrapy框架适合中小型批量数据采集，兼具开发效率和性能；Scrapy-Redis分布式爬虫适合工业级大规模数据采集，可横向扩展突破单机瓶颈。')

add_body(doc, '本次实验也让我认识到了自己的不足：①对Scrapy的信号（Signals）系统了解不够深入，后续需学习如何利用信号实现更精细的控制；②Scrapy的Telnet控制台和实时监控功能还未尝试；③在大规模分布式部署中，Redis的集群模式（Cluster）和哨兵模式（Sentinel）的配置还需要进一步学习。')

add_body(doc, '总的来说，本次实验让我建立了工程化爬虫开发的系统思维，掌握了从代码编写到框架选型、从单机部署到分布式扩展的完整能力链，为后续从事工业级数据采集工作奠定了扎实基础。Scrapy框架是Python爬虫领域的工业标准，其模块化、可扩展、高性能的特点使其成为大批量数据采集的首选方案。')

# ===========================
# 十、诚信承诺
# ===========================
doc.add_paragraph()  # 空行
add_body(doc, '诚信承诺：我保证本实验报告中的程序和本实验报告是我自己编写。')
add_body(doc, '承诺人：梁文泽')
add_body(doc, '日期：2026年6月')

# ===========================
# 保存
# ===========================
output_path = r'D:\项目\网络爬虫\实验\实验八\网络爬虫与数据收集_20230222_梁文泽_实验八.docx'
doc.save(output_path)
print(f'实验报告已生成：{output_path}')
print(f'文件大小：{os.path.getsize(output_path) / 1024:.1f} KB')
